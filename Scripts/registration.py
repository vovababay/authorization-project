from tkinter import *
from tkinter import messagebox
import pickle
from RegistrAccount import reg
import datetime
from AttemptsAccount import AddKeyMessenger
import getpass
import easygui
from Registrachui import import_photo_bd
import  docx

path_log_file='Logs\\registration.docx'


doc = docx.Document()
doc.add_paragraph(f"\n")
doc.add_paragraph("---------------------------------------------------------------------------------------------------------------\n")
doc.add_paragraph("---------------------------------------------------------------------------------------------------------------\n")
doc.add_paragraph(f"{datetime.datetime.now()} The program registration is running\n")
doc.save(path_log_file)




root=Tk()
root.geometry("380x500")
root.title("Регистрация в системе")
root.resizable(width=False,height=False)



class Photo():
    

    def __init__(self):
        self.path_photo=''

    def findfile(self):
            input_file = easygui.fileopenbox(default=f"C:\\Users\\{getpass.getuser()}\\Pictures\\",filetypes=["*.jpeg","*.jpg","*.png"])
            self.path_photo=input_file

def registration():
    
    
    
    
    text=Label(text="Для входа в систему- зарегестрируйтесь")
    text_Name = Label(text="Введите имя")
    registr_Name = Entry()
    text_login=Label(text="Введите логин:")
    registr_login=Entry()
    text_password=Label(text="Введите пароль:")
    registr_password=Entry(show="*")
    text_password1=Label(text="Повторите пароль:")
    registr_password1=Entry(show="*")
    button_registr=Button(text="Зарегестрироватся", command=lambda: save())
    text.pack()

    text_Name.pack()
    registr_Name.pack()

    text_login.pack()
    registr_login.pack()

    text_password.pack()
    registr_password.pack()

    text_password1.pack()
    registr_password1.pack()
    btn_path_text=Label(text="Выберите фотографию")
    btn_path_text.pack()

    
    
    
    btn_path=Button(text="Поиск фото",command=lambda: photo_user.findfile())
    btn_path.pack()
    text_variable=Label(text="Выбор мессенджера")
    text_variable.pack()
    variable = StringVar(root)
    variable.set("Telegram") # default value

    w = OptionMenu(root, variable, "Telegram")
    w.pack()

   
    button_registr.pack()
    photo_user=Photo()

    def save():
        if photo_user.path_photo == "":
            messagebox.showerror("Ошибка фотографии", "Проверьте что вы точно загружаете фотографию")
        else:
            is_photo = import_photo_bd(photo_user.path_photo, registr_login.get())


            doc = docx.Document(path_log_file)
            doc.add_paragraph(f"{datetime.datetime.now()} The registration window is open\n")
            doc.save(path_log_file)
            if registr_password.get() == registr_password1.get():
                is_reg = reg(registr_login.get(), registr_password.get(), variable.get(), is_photo, registr_Name.get())
                if is_reg == "Успешно":
                    messagebox.showinfo("Успешная регистрация",
                                        f"Вы успешно зарегестрировались в системе, ваш ключ авторизации в месседжере \"{AddKeyMessenger(registr_login.get())}\"")
                    # clear_registration_form()


                    doc = docx.Document(path_log_file)
                    doc.add_paragraph(f"{datetime.datetime.now()} Registered user {registr_login.get()}\n")
                    doc.save(path_log_file)
                elif is_reg == "Аккаунт существует":
                    messagebox.showerror("Ошибка регистрации", "Такой логин уже существует")
                    doc = docx.Document(path_log_file)
                    doc.add_paragraph(f"{datetime.datetime.now()} Registry errors, the user {registr_login.get()} exists\n")
                    doc.save(path_log_file)
                elif is_reg == "Пустой ввод":
                    messagebox.showerror("Ошибка регистрации", "Логин или пароль не могут быть пустыми")
                    doc = docx.Document(path_log_file)
                    doc.add_paragraph(f"{datetime.datetime.now()} The registration failed, the username or password is empty\n")
                    doc.save(path_log_file)
            else:
                messagebox.showerror("Ошибка регистрации", "Пароли не совпадают")

                doc = docx.Document(path_log_file)
                doc.add_paragraph(f"{datetime.datetime.now()} Password mismatch\n")
                doc.save(path_log_file)


registration()
root.mainloop()



'''
def clear_registration_form():
        text.pack_forget()

        text_login.pack_forget()
        registr_login.pack_forget()

        text_password.pack_forget()
        registr_password.pack_forget()

        text_password1.pack_forget()
        registr_password1.pack_forget()

        button_registr.pack_forget()

'''