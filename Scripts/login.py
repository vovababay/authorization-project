from tkinter import *
from tkinter import messagebox
from AttemptsAccount import Authentify
import datetime
from CodeCheck import Check
from CreatePass import CreatePassword
from IsTodayLoginMessenger import ChekTodayLoginInMessenger
import LockAccount
from MessageTelegram import MessegeLock
import time
from multiprocessing import *
import schedule
import  docx
from telebot import  types


path_log_file='Logs\\login.docx'
doc = docx.Document()
doc.add_paragraph("\n")
doc.add_paragraph("---------------------------------------------------------------------------------------------------------------\n")
doc.add_paragraph("---------------------------------------------------------------------------------------------------------------\n")
doc.add_paragraph(f"{datetime.datetime.now()} The program login is running\n")

doc.save(path_log_file)
root=Tk()
root.geometry("380x500")
root.title("Войти в систему")
root.resizable(width=False,height=False)



def login():
    doc = docx.Document(path_log_file)
    doc.add_paragraph(f"{datetime.datetime.now()} The login window is open\n")
    doc.save(path_log_file)

    text_enter_login=Label(text="Введите ваш логин:")
    enter_login=Entry()
    text_enter_password=Label(text="Введите ваш пароль:")
    enter_password=Entry(show="*")
    button_enter=Button(text="Войти",command=lambda: log_pass())

    
    
    text_enter_login.pack()
    
    enter_login.pack()
    
    text_enter_password.pack()
    
    enter_password.pack()
    
    button_enter.pack()

    
    def clear_login_form():
        text_enter_login.pack_forget()
        enter_login.pack_forget()

        text_enter_password.pack_forget()
        enter_password.pack_forget()

        button_enter.pack_forget()



    def log_pass():


        temp_bool=Authentify(enter_login.get(),enter_password.get())




        if temp_bool == True:


            islock=(LockAccount.IsLockAccount(enter_login.get()))
            if islock==False:


                MessegeLock(enter_login.get(), "На ваш аккаунт пытаются войти.")

                if ChekTodayLoginInMessenger(enter_login.get())==False:
                    messagebox.showerror("Ошибка","Авторизуйтесь в телеграм боте и повторите авторизацию заного")



                else:
                    messagebox.showinfo("Успешно","Авторизация через логин и пароль прошла успешно")
                    doc = docx.Document(path_log_file)
                    doc.add_paragraph(f"{datetime.datetime.now()} Successful login to your account '{enter_login.get()}'\n")
                    doc.save(path_log_file)
                    clear_login_form()
                    Check(root,enter_login.get())
            else:
                messagebox.showerror("Ошибка","Аккаунт заблокирован")

        else:
            messagebox.showerror("Ошибка","Неверный логин или пароль")
            doc = docx.Document(path_log_file)
            doc.add_paragraph(f"{datetime.datetime.now()} Unsuccessful login to your account '{enter_login.get()}'\n")
            doc.save(path_log_file)

                

        




login()

root.mainloop()




